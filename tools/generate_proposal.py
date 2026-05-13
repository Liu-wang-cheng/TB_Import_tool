"""基于模板生成 AI 应用大赛方案书"""

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy

TEMPLATE = 'C:/Users/LD/Desktop/AI应用大赛_方案书模板.docx'
OUTPUT = 'C:/Users/LD/Desktop/HCT智能缺陷管理平台_AI应用大赛方案书.docx'

doc = Document(TEMPLATE)

# ── 辅助函数 ──

def replace_paragraph_text(para, new_text, keep_style=True):
    """替换段落文本，保留第一个 run 的样式"""
    if not para.runs:
        para.text = new_text
        return
    # 保留第一个 run 的格式
    first_run = para.runs[0]
    first_run.text = new_text
    # 删除多余 run
    for r in para.runs[1:]:
        r.text = ''


def replace_cell_text(cell, new_text):
    """替换表格单元格文本，保留样式"""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
        for p in cell.paragraphs[1:]:
            for r in p.runs:
                r.text = ''
    else:
        cell.text = new_text


def fill_table(table, data):
    """用二维列表填充表格（跳过表头），不足时自动添加行"""
    existing = len(table.rows) - 1  # 减去表头
    for ri, row_data in enumerate(data):
        if ri < existing:
            row = table.rows[ri + 1]
        else:
            row = table.add_row()
        for ci, val in enumerate(row_data):
            replace_cell_text(row.cells[ci], val)


# ── 封面页 ──

replace_paragraph_text(doc.paragraphs[0], 'HCT 智能缺陷管理平台')
replace_paragraph_text(doc.paragraphs[5], '禅道→Teambition 智能 Bug 同步与分类平台')

# ── 提交日期 ──

replace_paragraph_text(doc.paragraphs[13], '提交日期：2026年05月08日')

# ── 封面表格（团队信息）──

table0 = doc.tables[0]
replace_cell_text(table0.rows[0].cells[1], '智能缺陷管理小组')
replace_cell_text(table0.rows[1].cells[1], '软件测试部')
replace_cell_text(table0.rows[2].cells[1], '胡继珍、刘旺诚')
replace_cell_text(table0.rows[3].cells[1], '胡继珍')
replace_cell_text(table0.rows[4].cells[1], 'hujizhen@hctrobot.com')
replace_cell_text(table0.rows[5].cells[1], 'V1.0')

# ── 1.1 方案名称 ──

replace_paragraph_text(doc.paragraphs[17],
    'HCT 智能缺陷管理平台：基于 AI 的禅道→Teambition Bug 自动同步与智能分类系统')

# ── 1.2 一句话描述 ──

replace_paragraph_text(doc.paragraphs[20],
    '基于 TF-IDF + LLM 多层管道的缺陷自动同步与智能分类平台，'
    '实现禅道 Bug 一键同步到 Teambition 并自动完成 SN 提取、缺陷分类、附件迁移，'
    '将单条 Bug 处理时间从 15 分钟压缩至 30 秒。')

# ── 1.3 方案摘要 ──

replace_paragraph_text(doc.paragraphs[29],
    '扫地机器人产品线每天产生大量固件缺陷，测试工程师需手动将禅道 Bug 逐条复制到 Teambition 进行项目管理，'
    '涉及标题翻译、严重程度映射、SN 编码提取、缺陷分类、附件上传等繁琐操作，每条耗时 15 分钟以上。'
    '本方案构建了"TF-IDF 本地学习 + LLM 大模型 + 规则兜底"三层 AI 分类管道，'
    '配合智能 SN/时间提取、双向标题同步、附件 OSS 直传等功能，'
    '实现从禅道到 Teambition 的端到端自动化同步，准确率超 95%，日均节省测试团队 3 小时以上。')

# ── 1.4 参赛类别 ──

# 替换参赛类别选项
replace_paragraph_text(doc.paragraphs[33], '☑ 效率提升类（流程自动化、信息检索、文档生成等）')
replace_paragraph_text(doc.paragraphs[37], '☑ 研发赋能类（代码生成、测试辅助、技术文档等）')

# ── 2.1 业务背景 ──

replace_paragraph_text(doc.paragraphs[47],
    '我司专注于家用扫地机器人（HS341 等型号）的研发与测试，固件迭代周期短（1-2 周），'
    '每次版本发布后产生大量缺陷报告。缺陷管理流程涉及两个系统：'
    '禅道（研发团队使用）和 Teambition（项目管理使用）。'
    '测试工程师是主要操作者，需要将禅道的每一条 Bug 手动搬运到 Teambition，'
    '包括翻译标题格式、映射严重程度、填入 SN 编码和缺陷产生时间、选择缺陷分类、上传附件等。'
    '涉及的角色包括：测试工程师（日常操作）、项目经理（跟踪进度）、研发工程师（修复缺陷）、'
    '产品经理（需求对齐）。')

# ── 2.2 核心痛点 ──

fill_table(doc.tables[1], [
    ['P1', '手动搬运 Bug 耗时巨大', '每条 Bug 需 15-20 分钟，20 条/天需 5+ 小时', '日均耗时 5 小时'],
    ['P2', '缺陷分类依赖人工经验', '分类标准不统一，新员工难以准确分类', '分类错误率约 15%'],
    ['P3', 'SN 编码和时间需手动提取', '从重现步骤文本中找 SN 和时间，容易遗漏或抄错', '遗漏率约 30%'],
    ['P4', '附件同步困难', '禅道图片和文件需下载后重新上传，流程繁琐', '每条含图片的 Bug 多花 5 分钟'],
])

# ── 2.3 现有方案的不足 ──

replace_paragraph_text(doc.paragraphs[101],
    '现有方案局限 1：纯人工操作，效率低且容易遗漏字段，重复劳动占据测试工程师大量时间')
replace_paragraph_text(doc.paragraphs[102],
    '现有方案局限 2：缺陷分类缺乏统一标准，依赖个人经验，新人上手成本高，分类一致性差')
replace_paragraph_text(doc.paragraphs[103],
    '现有方案局限 3：SN 编码和缺陷产生时间分散在重现步骤文本中，无结构化字段，人工提取耗时易错')

# ── 3.1 解决思路 ──

replace_paragraph_text(doc.paragraphs[66],
    '采用"本地学习 + 云端 AI + 规则兜底"三层管道架构：\n'
    '第一层：TF-IDF 相似度匹配（本地）— 从 Teambition 已有缺陷任务中深度学习分类模式，'
    '使用 jieba 中文分词 + scikit-learn TF-IDF + 余弦相似度，无需 GPU，无需外部 API，'
    '首次运行自动从 TB 拉取最新 5000 条缺陷任务训练模型，后续每 7 天增量学习。\n'
    '第二层：LLM 大模型分类 — 部署在内网的大语言模型（MiniMax-M2.7），对 TF-IDF 无法归类的缺陷进行 AI 分类，'
    '支持批量处理（每批 30 条），并设有 AI 审核机制抽检训练样本质量。\n'
    '第三层：规则兜底 — 按指派人部门前缀自动映射到对应分类（如 IOT-陈斌 → IOT-其他问题），'
    '确保 100% 有分类结果。\n\n'
    '与传统人工方案的差异：从"人工逐条处理"变为"AI 批量自动处理 + 人工抽检"，'
    '处理速度提升 30 倍，分类准确率从 85% 提升至 95%+。')

# ── 3.2 功能模块设计 ──

fill_table(doc.tables[2], [
    ['M1', '智能同步引擎', '禅道 Bug 自动同步到 Teambition：标题格式转换、严重程度映射、'
     '双向标题标注（禅道↔TB）、去重检测', '核心功能'],
    ['M2', 'AI 缺陷分类器', 'TF-IDF 本地学习 + LLM 大模型 + 规则兜底三层管道，'
     '自动将缺陷归入 40+ 个分类', '核心功能'],
    ['M3', '智能字段提取器', '从重现步骤文本中自动提取 SN 编码和缺陷产生时间，'
     '支持 M/D、纯时分等多种日期格式，SN 模式自动学习', '核心功能'],
    ['M4', '附件同步模块', '禅道附件和内联图片自动上传到 Teambition OSS，'
     '写入"日志附件"自定义字段，支持重试和大文件超时', '增强功能'],
    ['M5', 'PyQt6 GUI 界面', '可视化操作界面：筛选面板、进度条、日志区、配置管理，'
     '打包为 exe 供非技术人员使用', '增强功能'],
])

# ── 3.3 技术架构 ──

replace_paragraph_text(doc.paragraphs[79],
    '技术栈：Python 3.13 + PyQt6 + scikit-learn + jieba + requests\n\n'
    '数据来源：禅道 REST API v1（Bug 列表/详情/评论/附件）、Teambition Open API（JWT 认证、'
    '任务 CRUD、文件上传到阿里云 OSS）\n\n'
    'AI 模型：内网部署的 MiniMax-M2.7（通过 OpenAI 兼容 API 调用），'
    'TF-IDF 模型使用 scikit-learn 本地训练，缓存为 pickle 文件\n\n'
    '系统集成：禅道 REST API（Token + Session 双认证）、Teambition 企业应用（appToken 模式）、'
    '钉钉机器人 Webhook（同步结果通知）、阿里云 OSS STS 临时凭证上传\n\n'
    '安全措施：LLM 部署在内网（192.168.x.x），数据不经过公网；'
    '禅道和 TB 凭证存储在本地 YAML 配置文件；打包时自动剥离 API Key')

# ── 3.4 AI 工具与平台选型 ──

fill_table(doc.tables[3], [
    ['MiniMax-M2.7（内网部署）', '缺陷分类、SN/时间提取兜底、训练数据审核', '内网部署数据安全，OpenAI 兼容接口易集成'],
    ['scikit-learn TF-IDF', '本地学习缺陷分类模式，无需 GPU', '离线运行，延迟低，无 API 成本'],
    ['jieba 中文分词', '缺陷标题和描述的中文分词', '轻量高效，适合专业领域文本'],
])

# ── 4.1 现有工作流说明 ──

replace_paragraph_text(doc.paragraphs[87],
    '步骤 1：测试工程师在禅道中提交 Bug，填写标题、重现步骤、附件等')
replace_paragraph_text(doc.paragraphs[88],
    '步骤 2：打开 Teambition，手动创建任务，逐字段复制禅道内容（标题、严重程度、SN、版本等）')
replace_paragraph_text(doc.paragraphs[89],
    '步骤 3：手动下载禅道附件并上传到 Teambition 任务')
replace_paragraph_text(doc.paragraphs[90],
    '步骤 4：根据经验判断缺陷分类，手动选择分类标签')

# ── 4.2 AI 整合后的新工作流 ──

fill_table(doc.tables[4], [
    ['1', '手动创建任务，逐字段复制（15 分钟）', '点击"正式同步"按钮，AI 全自动处理（30 秒）'],
    ['2', '手动下载/上传附件（5 分钟）', '附件和内联图片自动上传到 OSS（自动）'],
    ['3', '凭经验手动分类（2 分钟）', 'TF-IDF + LLM 自动分类，准确率 95%+（自动）'],
    ['4', '手动从文本中找 SN 和时间（3 分钟）', '正则 + LLM 自动提取 SN 和缺陷时间（自动）'],
])

# ── 4.3 使用场景示例 ──

replace_paragraph_text(doc.paragraphs[99], '■ 场景 1：批量同步当日缺陷')
replace_paragraph_text(doc.paragraphs[100], '用户角色：测试工程师')
replace_paragraph_text(doc.paragraphs[101],
    '触发条件：每日下班前需将当天禅道新增 Bug 同步到 Teambition')
replace_paragraph_text(doc.paragraphs[102],
    '操作步骤：打开工具 → 选择日期范围和指派人筛选 → 点击"试运行"预览 → 确认后点击"正式同步" → '
    'AI 自动完成标题转换、SN 提取、缺陷分类、附件上传 → 钉钉自动推送同步结果')
replace_paragraph_text(doc.paragraphs[103],
    '预期效果：20 条 Bug 同步从原来 5 小时缩短至 10 分钟，包含附件和分类全部自动完成')

replace_paragraph_text(doc.paragraphs[105], '■ 场景 2：Bug 重新激活自动追踪')
replace_paragraph_text(doc.paragraphs[106], '用户角色：测试工程师')
replace_paragraph_text(doc.paragraphs[107],
    '触发条件：禅道中已关闭的 Bug 被开发重新激活')
replace_paragraph_text(doc.paragraphs[108],
    '操作步骤：工具检测到禅道 Bug active 但 TB 任务已关闭 → '
    '自动重新打开 TB 任务 → 更新执行人为当前指派人 → 同步最新评论和附件 → 添加重新激活评论')
replace_paragraph_text(doc.paragraphs[109],
    '预期效果：无需人工追踪 reopened 状态，避免遗漏回归 Bug')

# ── 5.1 量化收益 ──

fill_table(doc.tables[5], [
    ['时间效率', '每条 Bug 手动搬运 15-20 分钟', 'AI 自动同步约 30 秒/条', '↑ 97%'],
    ['分类准确率', '人工分类错误率约 15%', 'AI 分类准确率 95%+', '↑ 80%'],
    ['人力成本', '每日需 1 人花费 5 小时搬运', '每日一键同步，10 分钟完成', '↓ 95%'],
    ['附件同步', '每张图片需手动下载上传', '图片和附件自动 OSS 直传', '全自动'],
])

# ── 5.2 质性价值 ──

replace_paragraph_text(doc.paragraphs[118],
    '员工体验：测试工程师从繁琐的搬运工作中解放出来，聚焦于测试用例设计和质量分析')
replace_paragraph_text(doc.paragraphs[119],
    '决策质量：AI 分类基于历史数据学习，减少主观偏差，分类一致性大幅提升')
replace_paragraph_text(doc.paragraphs[120],
    '知识传承：TF-IDF 模型沉淀了团队的历史分类经验，新员工无需培训即可使用统一标准')
replace_paragraph_text(doc.paragraphs[121],
    '组织能力：推动缺陷管理流程标准化，为后续自动化测试平台建设奠定基础')

# ── 5.3 创新亮点 ──

replace_paragraph_text(doc.paragraphs[125],
    '创新点 1：独创"TF-IDF + LLM + 规则兜底"三层分类管道，'
    '本地学习无 API 成本，LLM 只处理 TF-IDF 无法归类的少量样本，大幅降低调用成本')
replace_paragraph_text(doc.paragraphs[126],
    '创新点 2：SN 编码自动学习机制——从 Teambition 已有任务的 SN 字段中学习项目特定的正则模式，'
    '适配不同产品线的 SN 格式差异，无需手动配置')
replace_paragraph_text(doc.paragraphs[127],
    '创新点 3：AI 审核训练数据——LLM 抽检 TF-IDF 训练样本，自动剔除不合理分类，'
    '确保模型质量持续提升')

# ── 6.1 可复用性分析 ──

fill_table(doc.tables[6], [
    ['缺陷分类管道', '低 ✅', 'TF-IDF + LLM 管道可适配任意缺陷分类体系，只需更换分类描述配置即可'],
    ['智能同步引擎', '低 ✅', '禅道→TB 同步逻辑可直接用于其他使用相同工具链的项目组'],
    ['SN 学习与提取', '中 ⚠️', '需各产品线从自己的 TB 任务中学习 SN 模式，约需 10 分钟自动完成'],
    ['附件 OSS 直传', '低 ✅', '基于 Teambition 标准 API，无需额外配置'],
    ['GUI 界面', '低 ✅', '已打包为 exe，双击即可使用，非技术人员零门槛'],
])

# ── 6.2 推广方案 ──

replace_paragraph_text(doc.paragraphs[138], '目标团队：HS341 产品线测试小组先行试用（5 人）')
replace_paragraph_text(doc.paragraphs[139], '验证指标：单条 Bug 同步时间、分类准确率、SN 提取覆盖率')
replace_paragraph_text(doc.paragraphs[140], '关键里程碑：完成 100 条真实 Bug 同步测试，准确率超 95%')

replace_paragraph_text(doc.paragraphs[143], '推广范围：扩展至全部产品线测试团队（约 30 人）')
replace_paragraph_text(doc.paragraphs[144], '配套措施：配置各产品线的禅道和 TB 参数，制作快速上手视频教程')
replace_paragraph_text(doc.paragraphs[145], '反馈机制：通过钉钉机器人收集使用反馈，每月迭代优化分类模型')

replace_paragraph_text(doc.paragraphs[148], '推广形式：作为测试团队标准工具，纳入新员工 onboarding 培训流程')
replace_paragraph_text(doc.paragraphs[149], '运营支持：维护 FAQ 文档，指定各产品线一位对接人')
replace_paragraph_text(doc.paragraphs[150], '成功标准：月活用户达 30 人，日均同步 100+ 条 Bug，用户满意度 4.5/5.0')

# ── 6.3 所需资源与支持 ──

replace_paragraph_text(doc.paragraphs[154],
    '技术支持：IT 部门协助配置 Teambition 企业应用权限和禅道 API 访问，预计 2 人天')
replace_paragraph_text(doc.paragraphs[155],
    '数据支持：开放各产品线的禅道和 Teambition 项目访问权限')
replace_paragraph_text(doc.paragraphs[156],
    '人力投入：1 名测试工程师兼职维护分类模型和配置，每周约 2 小时')
replace_paragraph_text(doc.paragraphs[157],
    '预算需求：内网 LLM 已部署，无额外 API 费用；阿里云 OSS 存储费用约 50 元/月')

# ── 7.1 实施路线图 ──

fill_table(doc.tables[7], [
    ['准备期', '第 1～2 周', '完善 AI 分类器、训练 TF-IDF 模型、配置各产品线参数', '工具 v1.2 exe、配置文件'],
    ['试点期', '第 3～4 周', 'HS341 产品线试用，每日同步验证，收集分类准确率数据', '试用报告、模型优化 V1.3'],
    ['推广期', '第 5～8 周', '扩展至全部产品线，制作培训材料，开展使用培训', '用户手册、培训记录、运营数据'],
    ['稳定运营', '第 9 周起', '持续运营、监控 KPI、定期增量学习、钉钉通知集成', '月度运营报告、迭代日志'],
])

# ── 7.2 风险分析与应对 ──

fill_table(doc.tables[8], [
    ['TF-IDF 冷启动（新项目无历史数据）', '中', '中', '新项目先用 LLM 和规则兜底分类，积累 50+ 条后自动训练 TF-IDF'],
    ['LLM 分类结果不稳定', '低', '低', '设置 temperature=0.1 降低随机性，增加人工抽检环节'],
    ['禅道或 TB API 限流', '低', '低', '内置批量暂停（每 20 条暂停 3 秒）和指数退避重试机制'],
    ['附件上传超时', '低', '低', '按文件大小动态计算超时（100KB/s 保底），最多重试 3 次'],
])

# ── 7.3 成功衡量标准（KPI）──

replace_paragraph_text(doc.paragraphs[166],
    'KPI 1：工具上线后 1 个月内，单条 Bug 同步时间从 15 分钟降至 30 秒以内')
replace_paragraph_text(doc.paragraphs[167],
    'KPI 2：AI 缺陷分类准确率稳定在 95% 以上（人工抽检 100 条验证）')
replace_paragraph_text(doc.paragraphs[168],
    'KPI 3：SN 编码和缺陷时间自动提取覆盖率达到 90% 以上')
replace_paragraph_text(doc.paragraphs[169],
    'KPI 4：月度有效反馈并迭代更新不少于 2 次，模型持续优化')

# ── 8.1 团队成员 ──

fill_table(doc.tables[9], [
    ['胡继珍（队长）', '软件测试部', '方案设计 / 项目统筹', '整体方案设计，AI 分类器开发，对外沟通协调'],
    ['刘旺诚', '软件测试部', 'AI 技术实现', '同步引擎开发、智能提取器实现、GUI 界面开发'],
])

# ── 8.2 相关经验与背景 ──

replace_paragraph_text(doc.paragraphs[176],
    '深耕扫地机器人测试领域多年，对固件、APP、导航等测试流程有深刻理解')
replace_paragraph_text(doc.paragraphs[177],
    '熟悉 Python 自动化测试（PyTest、Appium、Selenium），有丰富的工具开发经验')
replace_paragraph_text(doc.paragraphs[178],
    '对 AI 工具有深入实践，掌握 LLM 提示词工程、TF-IDF 文本分类、正则提取等技术')

# ── 8.3 为什么我们能做好这个方案 ──

replace_paragraph_text(doc.paragraphs[182],
    '我们深耕扫地机器人测试领域多年，对"禅道→Teambition"这条缺陷管理链路的痛点有切身体会——'
    '因为这套工具最初就是为了解决我们自己的日常工作痛点而开发的。'
    '从 v1.0 的基础同步到 v1.2 的 AI 智能分类和自动提取，每一步都经过真实业务场景验证。'
    '工具已在 HS341 产品线日常使用，日均处理 20+ 条 Bug，分类准确率持续提升。'
    '我们对 AI 工具有足够的热情和探索精神，能够快速迭代验证，让工具真正服务于业务。')

# ── 保存 ──

doc.save(OUTPUT)
print(f'方案书已生成: {OUTPUT}')
